"""Self-contained Knowledge Base Seeder (src/api/db/seed_kb.py).

Downloads actual KB and FAQ articles from MinIO, parses them, and inserts them
directly into the Postgres database. Fallbacks to mock seeding on local/dev envs.
"""

from __future__ import annotations

import asyncio
import hashlib
import io
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine
from transliterate import translit  # type: ignore[import-untyped]

from src.api.articles.models import Article
from src.api.categories.models import Category
from src.api.config import get_settings

# ---------------------------------------------------------------------------
# Seed pinning (ADR-0027)

SEED_VERSION = "2026-05-28"
SEED_BUCKET = "kb-seed"
SEED_PREFIX = f"articles/{SEED_VERSION}"

# Pinned sha256 для текущей seed-версии.
EXPECTED_SHA256: dict[str, str] = {
    "reHome_FAQ_топ15.docx": "e4d0834db83e12d705176ba65e201fe9bf118eceea80186dcc328bb7d093272b",
    "reHome_База_статей_120.docx": "3e9db4cb0385c44679fe9687861fd62569bb1f4032ddeee9849137887d3ac05f",
}

# Маппинг audience значений из .docx → ArticleAudience literal.
AUDIENCE_MAP = {
    "all": "all",
    "guest": "guest",
    "tenant": "tenant",
    "landlord": "landlord",
    "agent": "agent",
    "staff": "staff",
}
ACCESS_MAP = {"PUBLIC", "LOGGED", "AGENT", "STAFF"}

# Local mock data
CATEGORIES: list[dict[str, str]] = [
    {"slug": "arenda", "title": "Аренда жилья", "description": "Поиск, бронирование и заселение."},
    {"slug": "platezhi", "title": "Оплата и эскроу", "description": "Платежи, гарантии, возврат залога."},
    {"slug": "verifikatsiya", "title": "Верификация и KYC", "description": "Подтверждение личности и собственника."},
    {"slug": "dogovor", "title": "Договор найма", "description": "Условия, подписание и расторжение."},
]

ARTICLES: list[dict[str, Any]] = [
    {
        "slug": "kak-zabronirovat-kvartiru",
        "title": "Как забронировать квартиру",
        "category": "arenda",
        "summary": "Пошагово: от поиска объекта до подтверждённой брони.",
        "tags": ["бронирование", "аренда", "начало"],
        "body_markdown": (
            "## Как забронировать квартиру в reHome\n\n"
            "1. Найдите подходящий объект в каталоге и откройте его карточку.\n"
            "2. Нажмите **«Забронировать»** и выберите даты заселения.\n"
            "3. Пройдите проверку личности (KYC), если ещё не проходили.\n"
            "4. Дождитесь подтверждения от собственника.\n\n"
            "После подтверждения бронь переходит в статус активной, и вы можете "
            "перейти к подписанию договора найма и оплате."
        ),
    },
    {
        "slug": "kak-vernut-zalog",
        "title": "Как вернуть залог",
        "category": "platezhi",
        "summary": "Когда и как возвращается обеспечительный платёж после выезда.",
        "tags": ["залог", "возврат", "выезд", "оплата"],
        "body_markdown": (
            "## Возврат залога\n\n"
            "Обеспечительный платёж (залог) хранится на эскроу-счёте и "
            "возвращается после выезда при отсутствии претензий.\n\n"
            "**Порядок возврата:**\n\n"
            "1. Подпишите акт приёма-передачи при выезде.\n"
            "2. Собственник проверяет состояние квартиры.\n"
            "3. Если претензий нет — залог возвращается на вашу карту в течение "
            "нескольких рабочих дней.\n\n"
            "Если возник спор по состоянию квартиры, возврат залога "
            "приостанавливается до разбирательства."
        ),
    },
    {
        "slug": "chto-takoe-eskrou",
        "title": "Что такое эскроу и зачем оно нужно",
        "category": "platezhi",
        "summary": "Как reHome защищает деньги нанимателя и собственника.",
        "tags": ["эскроу", "безопасность", "оплата"],
        "body_markdown": (
            "## Эскроу в reHome\n\n"
            "Эскроу — это защищённый счёт, на котором деньги удерживаются до "
            "выполнения условий сделки. Наниматель вносит оплату, но собственник "
            "получает её только после заселения и подтверждения.\n\n"
            "Это защищает обе стороны: наниматель уверен, что деньги не уйдут до "
            "заселения, а собственник — что оплата гарантирована."
        ),
    },
    {
        "slug": "kak-projti-verifikaciyu",
        "title": "Как пройти верификацию личности",
        "category": "verifikatsiya",
        "summary": "Подтверждение личности через банк или оператора связи.",
        "tags": ["kyc", "верификация", "паспорт"],
        "body_markdown": (
            "## Верификация личности (KYC)\n\n"
            "Перед бронированием необходимо подтвердить личность. reHome "
            "поддерживает несколько способов:\n\n"
            "- через мобильного оператора (МТС);\n"
            "- через банк (Сбер, Т-Банк);\n"
            "- вводом паспортных данных вручную (для собственников).\n\n"
            "Проверка занимает от нескольких секунд до пары минут. После успешной "
            "верификации ваш профиль получает статус «подтверждён»."
        ),
    },
    {
        "slug": "kak-podpisat-dogovor",
        "title": "Как подписать договор найма",
        "category": "dogovor",
        "summary": "Электронная подпись договора по SMS-коду.",
        "tags": ["договор", "подпись", "sms"],
        "body_markdown": (
            "## Подписание договора найма\n\n"
            "Договор найма подписывается электронно:\n\n"
            "1. Откройте сформированный договор в разделе сделки.\n"
            "2. Внимательно проверьте условия (сроки, сумма, адрес).\n"
            "3. Нажмите **«Подписать»** — на ваш телефон придёт SMS-код.\n"
            "4. Введите код для подтверждения подписи.\n\n"
            "Обе стороны (наниматель и собственник) подписывают договор своими "
            "кодами. После этого договор считается заключённым."
        ),
    },
    {
        "slug": "kak-rastorgnut-dogovor",
        "title": "Как досрочно расторгнуть договор",
        "category": "dogovor",
        "summary": "Условия и порядок досрочного расторжения найма.",
        "tags": ["договор", "расторжение", "выезд"],
        "body_markdown": (
            "## Досрочное расторжение договора\n\n"
            "Досрочно расторгнуть договор можно по соглашению сторон или в "
            "случаях, предусмотренных договором.\n\n"
            "1. Уведомите вторую сторону через платформу заранее (срок указан "
            "в договоре).\n"
            "2. Согласуйте дату выезда и подпишите акт приёма-передачи.\n"
            "3. После проверки квартиры производится взаиморасчёт и возврат "
            "залога.\n\n"
            "Если согласия достичь не удаётся, обращение передаётся в поддержку "
            "для разбирательства."
        ),
    },
]


# ---------------------------------------------------------------------------
# S3 Fetching

def _fetch_s3(bucket: str, key: str) -> bytes:
    """Получает объект из MinIO (или любого S3-compatible) по env-credentials."""
    from src.api.documents.storage import get_minio_client

    settings = get_settings()
    client = get_minio_client(settings)
    response = client.get_object(bucket_name=bucket, object_name=key)
    try:
        return bytes(response.read())
    finally:
        response.close()
        response.release_conn()


def fetch_source(uri: str) -> tuple[bytes, str]:
    """Загружает .docx по URI; возвращает (bytes, basename)."""
    parsed = urlparse(uri)
    scheme = parsed.scheme.lower()

    if scheme in ("", "file"):
        path = Path(parsed.path if scheme == "file" else uri)
        if not path.is_absolute():
            path = path.resolve()
        return path.read_bytes(), path.name

    if scheme == "seed":
        name = (parsed.netloc + parsed.path).lstrip("/")
        if not name:
            raise ValueError(f"seed:// URI без имени: {uri!r}")
        key = f"{SEED_PREFIX}/{name}"
        return _fetch_s3(SEED_BUCKET, key), name

    if scheme == "s3":
        bucket = parsed.netloc
        key = parsed.path.lstrip("/")
        if not bucket or not key:
            raise ValueError(f"s3:// URI должен быть s3://<bucket>/<key>: {uri!r}")
        return _fetch_s3(bucket, key), Path(key).name

    raise ValueError(
        f"Неизвестный URI scheme {scheme!r}; ожидаю file://, s3://, seed://, или абсолютный path"
    )


def verify_sha256(data: bytes, basename: str, *, skip: bool = False) -> None:
    """Проверяет sha256(data) совпадает с pinned значением для basename."""
    actual = hashlib.sha256(data).hexdigest()
    expected = EXPECTED_SHA256.get(basename)
    if skip or expected is None:
        if expected is None:
            print(f"  [sha256] {basename}: {actual} (no pinned hash — skip)")
        else:
            print(f"  [sha256] {basename}: {actual} (skip-verify)")
        return
    if actual != expected:
        raise ValueError(
            f"sha256 mismatch для {basename}:\n"
            f"  expected: {expected}\n"
            f"  actual:   {actual}"
        )
    print(f"  [sha256] {basename}: ok")


# ---------------------------------------------------------------------------
# Parsing

def to_slug(title: str, max_len: int = 80) -> str:
    """Cyrillic title → kebab-case latin slug."""
    try:
        s = translit(title, "ru", reversed=True)
    except Exception:
        s = title
    s = s.lower()
    s = re.sub(r"[^a-z0-9-]+", "-", s)
    s = re.sub(r"-+", "-", s).strip("-")
    return s[:max_len].rstrip("-") or "untitled"


def parse_tags(line: str) -> list[str]:
    """Парсит `tags: [«a», «b», «c»]` → ['a', 'b', 'c']."""
    line = re.sub(r"^[^:]+:\s*", "", line)
    line = line.strip("[]").strip()
    tags = []
    for chunk in re.split(r"[«»\"',]", line):
        chunk = chunk.strip()
        if chunk and len(chunk) <= 64:
            tags.append(chunk)
    return tags[:10]


def field_value(line: str, key: str) -> str:
    """Вытаскивает `key: value` из `Compact` paragraph."""
    pattern = rf"^{key}\s*:\s*(.*)$"
    m = re.match(pattern, line.strip(), flags=re.IGNORECASE)
    return m.group(1).strip() if m else ""


def parse_faq(data: bytes) -> list[dict[str, Any]]:
    """FAQ doc — каждая статья начинается с Heading 2 'FAQ-NNN'."""
    doc = Document(io.BytesIO(data))
    articles: list[dict[str, Any]] = []
    cur: dict[str, Any] | None = None
    body_lines: list[str] = []

    def flush() -> None:
        if cur is None:
            return
        cur["body_markdown"] = "\n\n".join(body_lines).strip()
        if cur.get("title") and cur["body_markdown"]:
            articles.append(cur)

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name == "Heading 2" and text.startswith("FAQ-"):
            flush()
            cur = {
                "category": "FAQ",
                "audience": "all",
                "access_level": "PUBLIC",
                "tags": ["topfaq"],
                "status": "PUBLISHED",
                "language": "ru",
            }
            body_lines = []
            continue
        if cur is None:
            continue
        if v := field_value(text, "question"):
            cur["title"] = v[:200]
            continue
        if v := field_value(text, "category"):
            cur["category"] = v[:100]
            continue
        if v := field_value(text, "audience"):
            cur["audience"] = AUDIENCE_MAP.get(v.lower(), "all")
            continue
        if v := field_value(text, "access_level"):
            v = v.upper()
            cur["access_level"] = v if v in ACCESS_MAP else "PUBLIC"
            continue
        if text.lower().startswith("tags"):
            parsed = parse_tags(text)
            seed = [t for t in cur.get("tags", []) if t not in parsed]
            cur["tags"] = (seed + parsed)[:10]
            continue
        if v := field_value(text, "short_answer"):
            body_lines.append(f"**Кратко:** {v}")
            continue
        if v := field_value(text, "full_answer"):
            body_lines.append(v)
            continue
        body_lines.append(text)
    flush()
    return articles


def parse_kb(data: bytes) -> list[dict[str, Any]]:
    """KB doc — категории как Heading 1, статьи как Heading 2 'Статья N'."""
    doc = Document(io.BytesIO(data))
    articles: list[dict[str, Any]] = []
    current_category = "Общее"
    cur: dict[str, Any] | None = None
    body_lines: list[str] = []

    def flush() -> None:
        if cur is None:
            return
        cur["body_markdown"] = "\n\n".join(body_lines).strip()
        if cur.get("title") and cur["body_markdown"]:
            articles.append(cur)

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        if p.style.name == "Heading 1" and text.startswith("Категория"):
            cat = re.sub(r"^Категория\s*\d+\.\s*", "", text).strip()
            current_category = cat[:100] or current_category
            flush()
            cur = None
            body_lines = []
            continue

        if p.style.name == "Heading 2" and text.startswith("Статья"):
            flush()
            cur = {
                "category": current_category,
                "audience": "all",
                "access_level": "PUBLIC",
                "tags": [],
                "status": "PUBLISHED",
                "language": "ru",
            }
            body_lines = []
            continue
        if cur is None:
            continue
        if v := field_value(text, "question"):
            cur["title"] = v[:200]
            continue
        if v := field_value(text, "audience"):
            cur["audience"] = AUDIENCE_MAP.get(v.lower(), "all")
            continue
        if v := field_value(text, "access_level"):
            v = v.upper()
            cur["access_level"] = v if v in ACCESS_MAP else "PUBLIC"
            continue
        if text.lower().startswith("tags"):
            cur["tags"] = parse_tags(text)
            continue
        body_lines.append(text)
    flush()
    return articles


# ---------------------------------------------------------------------------
# Seeding Main Logic

async def main() -> int:
    settings = get_settings()
    db_url = os.environ.get("DATABASE_URL") or settings.database_url
    engine = create_async_engine(db_url)
    factory = async_sessionmaker(engine, expire_on_commit=False)
    created_cats = created_arts = skipped = 0
    now = datetime.now(timezone.utc)

    env = os.environ.get("REHOME_ENV", settings.environment).lower()

    try:
        async with factory() as session:
            if env in ("prod", "staging"):
                print(f"Running ACTUAL KB seed for environment: {env}")
                try:
                    # 1. Fetch and parse actual articles from S3/MinIO
                    faq_bytes, faq_name = fetch_source("seed://reHome_FAQ_топ15.docx")
                    kb_bytes, kb_name = fetch_source("seed://reHome_База_статей_120.docx")

                    verify_sha256(faq_bytes, faq_name)
                    verify_sha256(kb_bytes, kb_name)

                    faq_articles = parse_faq(faq_bytes)
                    kb_articles = parse_kb(kb_bytes)
                    all_articles = faq_articles + kb_articles

                    # 2. Extract and insert unique categories
                    existing_cats = set(
                        (await session.execute(select(Category.slug))).scalars().all()
                    )
                    unique_categories = set(a["category"] for a in all_articles if a.get("category"))
                    for cat_name in unique_categories:
                        if cat_name in existing_cats:
                            continue
                        session.add(
                            Category(
                                slug=cat_name,
                                title=cat_name,
                                description=f"Статьи из категории {cat_name}",
                            )
                        )
                        created_cats += 1
                        existing_cats.add(cat_name)

                    # 3. Generate stable slugs and insert articles
                    existing_arts = set(
                        (await session.execute(select(Article.slug))).scalars().all()
                    )
                    for a in all_articles:
                        base_slug = to_slug(a["title"])
                        slug = base_slug
                        suffix = 1
                        while slug in existing_arts:
                            suffix += 1
                            slug = f"{base_slug}-{suffix}"[:80].rstrip("-")

                        if slug in existing_arts:
                            skipped += 1
                            continue

                        session.add(
                            Article(
                                slug=slug,
                                title=a["title"],
                                summary=a.get("summary", ""),
                                body_markdown=a["body_markdown"],
                                audience=a.get("audience", "all"),
                                language=a.get("language", "ru"),
                                category=a.get("category", "Общее"),
                                tags=a.get("tags", []),
                                access_level=a.get("access_level", "PUBLIC"),
                                status=a.get("status", "PUBLISHED"),
                                published_at=now,
                            )
                        )
                        created_arts += 1
                        existing_arts.add(slug)

                    await session.commit()
                    print(
                        f"OK: categories_created={created_cats}, articles_created={created_arts}, "
                        f"articles_skipped={skipped}"
                    )
                except Exception as exc:
                    print(f"FAILED to seed actual articles: {exc}")
                    raise exc
            else:
                # Local dev mock seeding
                print(f"Running MOCK KB seed for environment: {env}")
                existing_cats = set(
                    (await session.execute(select(Category.slug))).scalars().all()
                )
                for c in CATEGORIES:
                    if c["slug"] in existing_cats:
                        continue
                    session.add(
                        Category(
                            slug=c["slug"],
                            title=c["title"],
                            description=c["description"],
                        )
                    )
                    created_cats += 1

                existing_arts = set(
                    (await session.execute(select(Article.slug))).scalars().all()
                )
                for a in ARTICLES:
                    if a["slug"] in existing_arts:
                        skipped += 1
                        continue
                    session.add(
                        Article(
                            slug=a["slug"],
                            title=a["title"],
                            summary=a["summary"],
                            body_markdown=a["body_markdown"],
                            audience="all",
                            language="ru",
                            category=a["category"],
                            tags=a["tags"],
                            access_level="PUBLIC",
                            status="PUBLISHED",
                            published_at=now,
                        )
                    )
                    created_arts += 1

                await session.commit()
                print(
                    f"OK: categories_created={created_cats}, articles_created={created_arts}, "
                    f"articles_skipped={skipped}"
                )
    finally:
        await engine.dispose()
    return 0


if __name__ == "__main__":
    sys.exit(asyncio.run(main()))

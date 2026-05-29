"""Импортер статей из .docx → POST /api/v1/articles.

Парсит два .docx:
1. reHome_FAQ_топ15.docx — 15 FAQ (audience/access_level/tags явно
   указаны).
2. reHome_База_статей_120.docx — 120 KB-статей, organised в 11 категорий.
   Поля q/tags/audience указаны в paragraphs; access_level default PUBLIC.

Slug — transliterate(title) + kebab-case, suffix N если коллизия.
Все статусы → PUBLISHED.

Source resolution (ADR-0027 — seed source-of-truth = MinIO bucket
`kb-seed`). CLI принимает URI для каждого .docx:

    file:///abs/path/file.docx  — локальный (legacy / dev re-import)
    /abs/path/file.docx          — alias of file://
    s3://<bucket>/<key>          — MinIO / любой S3-compatible
    seed://<name>.docx           — alias of s3://kb-seed/articles/<DATE>/<name>.docx
                                   (DATE pinned в SEED_VERSION; см. README)

Defaults — pinned seed-версия 2026-05-28 в MinIO. Каждый источник имеет
expected sha256; reproducibility verified перед parse.
"""

from __future__ import annotations

import argparse
import hashlib
import io
import re
import sys
import time
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

import httpx
from docx import Document
from transliterate import translit  # type: ignore[import-untyped]

# ---------------------------------------------------------------------------
# Seed pinning (ADR-0027)

SEED_VERSION = "2026-05-28"
SEED_BUCKET = "kb-seed"
SEED_PREFIX = f"articles/{SEED_VERSION}"

# Pinned sha256 для текущей seed-версии. Обновляется при загрузке новой
# версии в MinIO + bump SEED_VERSION (см. backend/scripts/seed/README.md).
EXPECTED_SHA256: dict[str, str] = {
    "reHome_FAQ_топ15.docx": ("e4d0834db83e12d705176ba65e201fe9bf118eceea80186dcc328bb7d093272b"),
    "reHome_База_статей_120.docx": (
        "3e9db4cb0385c44679fe9687861fd62569bb1f4032ddeee9849137887d3ac05f"
    ),
}

API = "http://localhost:8000/api/v1"
TOKEN_PATH = Path("/tmp/.kb-token")

# Маппинг audience значений из .docx → ArticleAudience literal.
AUDIENCE_MAP = {
    "all": "all",
    "guest": "guest",
    "tenant": "tenant",
    "landlord": "landlord",
    "agent": "agent",
    "staff": "staff",
}
ACCESS_MAP = {"PUBLIC", "LOGGED", "AGENT", "STAFF"}


# ---------------------------------------------------------------------------
# Source fetch (ADR-0027)


def _fetch_s3(bucket: str, key: str) -> bytes:
    """Получает объект из MinIO (или любого S3-compatible) по env-credentials.

    Reuses `src.api.documents.storage.get_minio_client` — single source of
    truth для S3 config (endpoint, access_key, secret_key, secure).
    """
    # Ленивый импорт: основной API не должен тащить minio client при boot
    # без необходимости (ADR-0012). CLI вызывается отдельным процессом,
    # импорт здесь допустим.
    from src.api.config import get_settings
    from src.api.documents.storage import get_minio_client

    settings = get_settings()
    client = get_minio_client(settings)
    response = client.get_object(bucket_name=bucket, object_name=key)
    try:
        return bytes(response.read())
    finally:
        response.close()
        response.release_conn()


def fetch_source(uri: str) -> tuple[bytes, str]:
    """Загружает .docx по URI; возвращает (bytes, basename).

    basename используется для lookup'а в EXPECTED_SHA256.
    """
    parsed = urlparse(uri)
    scheme = parsed.scheme.lower()

    if scheme in ("", "file"):
        path = Path(parsed.path if scheme == "file" else uri)
        if not path.is_absolute():
            path = path.resolve()
        return path.read_bytes(), path.name

    if scheme == "seed":
        # `seed://<name>.docx` → s3://kb-seed/articles/<DATE>/<name>.docx
        name = (parsed.netloc + parsed.path).lstrip("/")
        if not name:
            raise ValueError(f"seed:// URI без имени: {uri!r}")
        key = f"{SEED_PREFIX}/{name}"
        return _fetch_s3(SEED_BUCKET, key), name

    if scheme == "s3":
        bucket = parsed.netloc
        key = parsed.path.lstrip("/")
        if not bucket or not key:
            raise ValueError(f"s3:// URI должен быть s3://<bucket>/<key>: {uri!r}")
        return _fetch_s3(bucket, key), Path(key).name

    raise ValueError(
        f"Неизвестный URI scheme {scheme!r}; ожидаю file://, s3://, seed://, или абсолютный path"
    )


def verify_sha256(data: bytes, basename: str, *, skip: bool) -> None:
    """Проверяет sha256(data) совпадает с pinned значением для basename.

    Если `skip=True` — печатает actual hash (для bumping версии) и
    пропускает проверку.
    """
    actual = hashlib.sha256(data).hexdigest()
    expected = EXPECTED_SHA256.get(basename)
    if skip or expected is None:
        if expected is None:
            print(f"  [sha256] {basename}: {actual} (no pinned hash — skip)")
        else:
            print(f"  [sha256] {basename}: {actual} (skip-verify)")
        return
    if actual != expected:
        raise SystemExit(
            f"sha256 mismatch для {basename}:\n"
            f"  expected: {expected}\n"
            f"  actual:   {actual}\n"
            f"Если это новая версия — bump SEED_VERSION + EXPECTED_SHA256 "
            f"(ADR-0027) или запусти с --no-verify-sha."
        )
    print(f"  [sha256] {basename}: ok")


# ---------------------------------------------------------------------------
# Slug + field parsing (без изменений с pre-ADR-0027 версии)


def to_slug(title: str, max_len: int = 80) -> str:
    """Cyrillic title → kebab-case latin slug."""
    try:
        s = translit(title, "ru", reversed=True)
    except Exception:
        s = title
    s = s.lower()
    s = re.sub(r"[^a-z0-9-]+", "-", s)
    s = re.sub(r"-+", "-", s).strip("-")
    return s[:max_len].rstrip("-") or "untitled"


def parse_tags(line: str) -> list[str]:
    """Парсит `tags: [«a», «b», «c»]` → ['a', 'b', 'c']."""
    line = re.sub(r"^[^:]+:\s*", "", line)
    line = line.strip("[]").strip()
    tags = []
    for chunk in re.split(r"[«»\"',]", line):
        chunk = chunk.strip()
        if chunk and len(chunk) <= 64:
            tags.append(chunk)
    return tags[:10]


def field_value(line: str, key: str) -> str:
    """Вытаскивает `key: value` из `Compact` paragraph."""
    pattern = rf"^{key}\s*:\s*(.*)$"
    m = re.match(pattern, line.strip(), flags=re.IGNORECASE)
    return m.group(1).strip() if m else ""


def parse_faq(data: bytes) -> list[dict[str, Any]]:
    """FAQ doc — каждая статья начинается с Heading 2 'FAQ-NNN'.

    Fields в Compact paragraphs (question, category, audience, access_level,
    tags). Тело — short_answer + full_answer + последующие paragraphs.
    """
    doc = Document(io.BytesIO(data))
    articles: list[dict[str, Any]] = []
    cur: dict[str, Any] | None = None
    body_lines: list[str] = []

    def flush() -> None:
        if cur is None:
            return
        cur["body_markdown"] = "\n\n".join(body_lines).strip()
        if cur.get("title") and cur["body_markdown"]:
            articles.append(cur)

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name == "Heading 2" and text.startswith("FAQ-"):
            flush()
            cur = {
                "category": "FAQ",
                "audience": "all",
                "access_level": "PUBLIC",
                "tags": ["topfaq"],
                "status": "PUBLISHED",
                "language": "ru",
            }
            body_lines = []
            continue
        if cur is None:
            continue
        if v := field_value(text, "question"):
            cur["title"] = v[:200]
            continue
        if v := field_value(text, "category"):
            cur["category"] = v[:100]
            continue
        if v := field_value(text, "audience"):
            cur["audience"] = AUDIENCE_MAP.get(v.lower(), "all")
            continue
        if v := field_value(text, "access_level"):
            v = v.upper()
            cur["access_level"] = v if v in ACCESS_MAP else "PUBLIC"
            continue
        if text.lower().startswith("tags"):
            parsed = parse_tags(text)
            seed = [t for t in cur.get("tags", []) if t not in parsed]
            cur["tags"] = (seed + parsed)[:10]
            continue
        if v := field_value(text, "short_answer"):
            body_lines.append(f"**Кратко:** {v}")
            continue
        if v := field_value(text, "full_answer"):
            body_lines.append(v)
            continue
        body_lines.append(text)
    flush()
    return articles


def parse_kb(data: bytes) -> list[dict[str, Any]]:
    """KB doc — категории как Heading 1, статьи как Heading 2 'Статья N'.

    Поля внутри: question, tags, audience. access_level не указан — default
    PUBLIC. Тело — все paragraphs после field-блока.
    """
    doc = Document(io.BytesIO(data))
    articles: list[dict[str, Any]] = []
    current_category = "Общее"
    cur: dict[str, Any] | None = None
    body_lines: list[str] = []

    def flush() -> None:
        if cur is None:
            return
        cur["body_markdown"] = "\n\n".join(body_lines).strip()
        if cur.get("title") and cur["body_markdown"]:
            articles.append(cur)

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        if p.style.name == "Heading 1" and text.startswith("Категория"):
            cat = re.sub(r"^Категория\s*\d+\.\s*", "", text).strip()
            current_category = cat[:100] or current_category
            flush()
            cur = None
            body_lines = []
            continue

        if p.style.name == "Heading 2" and text.startswith("Статья"):
            flush()
            cur = {
                "category": current_category,
                "audience": "all",
                "access_level": "PUBLIC",
                "tags": [],
                "status": "PUBLISHED",
                "language": "ru",
            }
            body_lines = []
            continue
        if cur is None:
            continue
        if v := field_value(text, "question"):
            cur["title"] = v[:200]
            continue
        if v := field_value(text, "audience"):
            cur["audience"] = AUDIENCE_MAP.get(v.lower(), "all")
            continue
        if v := field_value(text, "access_level"):
            v = v.upper()
            cur["access_level"] = v if v in ACCESS_MAP else "PUBLIC"
            continue
        if text.lower().startswith("tags"):
            cur["tags"] = parse_tags(text)
            continue
        body_lines.append(text)
    flush()
    return articles


def post_article(
    client: httpx.Client,
    art: dict[str, Any],
    used_slugs: set[str],
    headers: dict[str, str],
) -> str:
    """POST /api/v1/articles + idempotency on slug collision."""
    base_slug = to_slug(art["title"])
    slug = base_slug
    suffix = 1
    while slug in used_slugs:
        suffix += 1
        slug = f"{base_slug}-{suffix}"[:80].rstrip("-")
    used_slugs.add(slug)

    payload = {**art, "slug": slug}
    resp = client.post(f"{API}/articles", json=payload, headers=headers)
    if resp.status_code == 409:
        used_slugs.add(slug)
        return f"DUP {slug}"
    if resp.status_code != 201:
        return f"ERR {resp.status_code} {slug}: {resp.text[:200]}"
    return f"OK  {slug}"


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__.split("\n", 1)[0])
    parser.add_argument(
        "--faq",
        default="seed://reHome_FAQ_топ15.docx",
        help="URI для FAQ .docx (file://, s3://, seed://, или абсолютный path)",
    )
    parser.add_argument(
        "--kb",
        default="seed://reHome_База_статей_120.docx",
        help="URI для KB .docx",
    )
    parser.add_argument(
        "--no-verify-sha",
        action="store_true",
        help="Пропустить sha256 verification (для bumping seed-версии)",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Распарсить и показать sample article, без POST",
    )
    args = parser.parse_args()

    print(f"Fetching FAQ ← {args.faq}")
    faq_bytes, faq_name = fetch_source(args.faq)
    verify_sha256(faq_bytes, faq_name, skip=args.no_verify_sha)

    print(f"Fetching KB  ← {args.kb}")
    kb_bytes, kb_name = fetch_source(args.kb)
    verify_sha256(kb_bytes, kb_name, skip=args.no_verify_sha)

    print("\nParsing FAQ...")
    faq = parse_faq(faq_bytes)
    print(f"  FAQ articles: {len(faq)}")

    print("Parsing KB...")
    kb = parse_kb(kb_bytes)
    print(f"  KB articles: {len(kb)}")

    all_articles = faq + kb
    print(f"\nTotal to import: {len(all_articles)}")
    if all_articles:
        sample = all_articles[0]
        print(f"Sample article 0: title={sample['title'][:60]!r}")
        print(f"                  category={sample['category']!r}")
        print(f"                  audience={sample['audience']!r}")
        print(f"                  access_level={sample['access_level']!r}")
        print(f"                  tags={sample['tags']}")
        print(f"                  body chars={len(sample['body_markdown'])}")
    print()

    if args.dry_run:
        return 0

    token = TOKEN_PATH.read_text().strip()
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

    used_slugs: set[str] = set()
    ok = 0
    fail = 0
    dup = 0
    with httpx.Client(timeout=30.0) as client:
        for i, art in enumerate(all_articles, 1):
            try:
                result = post_article(client, art, used_slugs, headers)
            except Exception as exc:
                result = f"EXC {type(exc).__name__}: {exc}"
            if result.startswith("OK"):
                ok += 1
            elif result.startswith("DUP"):
                dup += 1
            else:
                fail += 1
            if i % 10 == 0 or not result.startswith("OK"):
                print(f"  [{i:3d}/{len(all_articles)}] {result}")
            time.sleep(0.02)

    print(f"\n=== Summary: OK={ok}, DUP={dup}, FAIL={fail} ===")
    return 0 if fail == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
